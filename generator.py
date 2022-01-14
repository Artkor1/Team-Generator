from google.oauth2.service_account import Credentials
import gspread
from gspread import utils
import time
import random
import numpy as np
import winsound

#on pycharm: View ->  Tool Windows -> Python Packages -> python-docx -> install
#pip install python-docx
from docx import Document

#on pycharm: View ->  Tool Windows -> Python Packages -> gspread-formatting -> install
#pip install gspread-formatting
from gspread_formatting import *

scope = ['https://www.googleapis.com/auth/spreadsheets',
         'https://www.googleapis.com/auth/drive.file',
         "https://www.googleapis.com/auth/drive"]
creds = Credentials.from_service_account_file('genroom-2654a6209a4a.json', scopes=scope)
client = gspread.authorize(creds)

spreadsheet = client.open("team generator")
sheet1 = spreadsheet.get_worksheet(0)
sheet2 = spreadsheet.get_worksheet(1)
sheet3 = spreadsheet.get_worksheet(2)


print("File opened")


# set cell color in multiple cells
def setColor(worksheet, color, rowStart, columnStart, rowEnd, columnEnd):
    cellRange = utils.rowcol_to_a1(rowStart, columnStart) + ":" + utils.rowcol_to_a1(rowEnd, columnEnd)
    if color == "red":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 1.0,
                "green": 0.5,
                "blue": 0.5
            }
        })
    elif color == "blue":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 0.5,
                "green": 0.5,
                "blue": 1.0
            }
        })
    elif color == "white":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 1.0,
                "green": 1.0,
                "blue": 1.0
            }
        })


#remove players from the list if the lack: timestamp, name, nickname or type of game they play
def removeEmpty(players):
    i = 0
    while i < len(players):
        for j in range(4):
            if players[i][j] == "":
                del players[i]
                i = i - 1
                continue
        i = i + 1


#remove duplicate entries of players in the list
def removeDuplicate(players):
    i = 0
    while i < len(players):
        j = i + 1
        while j < len(players):
            if players[i][1] == players[j][1] or players[i][2] == players[j][2]:
                del players[j]
                j = j - 1
            j = j + 1
        i = i + 1


#divide players between LOL and CS
def prepareGames(players):
    for player in players:
        if player[3] == "LoL":
            if useName:
                playersLOL.append(player[1])
            else:
                playersLOL.append(player[2])
        if player[3] == "CS:GO":
            if useName:
                playersCS.append(player[1])
            else:
                playersCS.append(player[2])


#remove players if there are too many, for example: reduce from 19 to 10
def removeExcessivePlayers(players):
    if len(players) < 10:
        print("Impossible to form a team with ", len(players))
        return
    while len(players) % 10 != 0:
        del players[-1]


#Find number of the previous game
def findPreviousGame(worksheet):
    previousGameNumber = 0
    query = re.compile("GAME .*")
    cells = worksheet.findall(query)
    time.sleep(1)
    if len(cells) != 0:
        games = []
        for cell in cells:
            games.append(cell.value)
        games = sorted(games)
        previousGame = games[-1]
        previousGameNumber = int(previousGame[-1])
    return previousGameNumber


#find 1st empty column to write in
def findAvailableColumn(worksheet):
    column = 1
    while column <= 26:
        cell = worksheet.cell(1, column).value
        time.sleep(1)
        if cell is None:
            break
        column = column + 1
    return column


#find 1st empty row to write in
def findAvailableRow(worksheet):
    str_list = list(filter(None, worksheet.col_values(1)))
    time.sleep(1)
    return str(len(str_list)+1)


#shuffle teams and write results to spreadsheet
def generateGame(players, teamNumbers, game):
    random.shuffle(players)
    sheet = spreadsheet.worksheet(game)

    column = findAvailableColumn(sheet)

    for i in range(0, teamNumbers):
        if i % 2 == 0:
            previousGame = findPreviousGame(sheet)
            sheet.update_cell((i / 2) * 13 + 1, column, "GAME " + str(previousGame + 1))
            time.sleep(1)
        sheet.update_cell(i * 6 + 2 + i / 2, column, "TEAM " + str(i + 1))
        time.sleep(1)

        j = i * 5
        while j < i*5 + 5:
            if game == "LOL":
                roomsList[i].append(players[j])
            else:
                roomsList[i+teamsLOL].append(players[j])

            sheet.update_cell(j + 3 + i + i / 2, column, players[j])
            time.sleep(1)
            j = j + 1
        if i % 2 == 0:
            setColor(sheet, "red", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        else:
            setColor(sheet, "blue", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        time.sleep(1)
        if i % 2 != 0:
            print(game, "game", (previousGame+1), "generated successfully")


def generateGamev2(players, teamNumbers, game):
    #playsound('losu losu.mp3', False)
    random.shuffle(players)
    if game == "LOL":
        sheet = sheet2
    elif game == "CS":
        sheet = sheet3

    row = 1
    column = findAvailableColumn(sheet)

    for i in range(0, teamNumbers):
        array = np.empty(shape=[0, 1])
        if i % 2 == 0:
            previousGame = findPreviousGame(sheet)
            array = np.append(array, [["GAME " + str(previousGame + 1)]], axis=0)
        array = np.append(array, [["TEAM " + str(i + 1)]], axis=0)

        j = i * 5
        while j < i * 5 + 5:
            if game == "LOL":
                roomsList[i].append(players[j])
            else:
                roomsList[i+teamsLOL].append(players[j])
            array = np.append(array, [[players[j]]], axis=0)
            j = j + 1

        cellRange = utils.rowcol_to_a1(row, column)
        row = row + len(array.tolist())
        sheet.update(cellRange, array.tolist())
        time.sleep(1)

        if i % 2 == 0:
            setColor(sheet, "red", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        else:
            setColor(sheet, "blue", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        time.sleep(1)
        if i % 2 != 0:
            print(game, "game", (previousGame+1), "generated successfully")


#write distribution of players across rooms to file
def writeRooms(rooms):
    f = open("rooms.txt", "w")
    for i in range(0, len(rooms)):
        if i != 0:
            f.write("\n\n")
        f.write("Room " + str(i+1))

        if i < teamsLOL:
            f.write(" (LOL)")
        else:
            f.write("(CS)")

        for j in range(0, len(rooms[i])):
            if j % 5 == 0:
                f.write("\nGame " + str(int(j / 5 + 1)) + ":  ")
            f.write(rooms[i][j])
            if (j + 1) % 5 != 0:
                f.write(", ")
    f.close()
    print("Rooms distribution file generated successfully")


def writeRoomsDOCX(rooms):
    document = Document()
    document.add_heading('Rooms', 0)
    for i in range(0, len(rooms)):
        heading = document.add_heading('Room' + (str(i+1)))
        if i < teamsLOL:
            heading.add_run(' (LOL)')
        else:
            heading.add_run(' (CS)')
        for j in range(0, len(rooms[i])):
            if j % 5 == 0:
                paragraph = document.add_paragraph('Game ' + str(int(j / 5 + 1)) + ':  ')
            paragraph.add_run(rooms[i][j])
            if (j + 1) % 5 != 0:
                paragraph.add_run(', ')
    document.save('rooms.docx')

    print("Rooms distribution file DOCX generated successfully")


def prepareWorksheets():
    # clear worksheets
    spreadsheet.get_worksheet(1).clear()
    spreadsheet.get_worksheet(2).clear()

    # set column width
    set_column_width(sheet2, 'A:Z', 135)
    set_column_width(sheet3, 'A:Z', 135)

    # clear colors
    setColor(sheet2, "white", 1, 1, 40, 26)
    setColor(sheet3, "white", 1, 1, 40, 26)

    # set font size and bold
    sheet2.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet2.format('A2:Z2', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet2.format('A8:Z8', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet2.format('A14:Z14', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet2.format('A15:Z15', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet2.format('A21:Z21', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet2.format('A27:Z27', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet2.format('A28:Z28', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet2.format('A34:Z34', {'textFormat': {'bold': True, "fontSize": 12}})  # Team

    sheet3.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet3.format('A2:Z2', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet3.format('A8:Z8', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet3.format('A14:Z14', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet3.format('A15:Z15', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet3.format('A21:Z21', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet3.format('A27:Z27', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
    sheet3.format('A28:Z28', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
    sheet3.format('A34:Z34', {'textFormat': {'bold': True, "fontSize": 12}})  # Team

    # this is the amount of previous operations
    # time.sleep(24)

    print("Worksheets prepared")


#play sound
winsound.PlaySound('sounds/losowanko zapraszam.wav', winsound.SND_ASYNC)
prepareWorksheets()

#flag to use either names or nicknames
useName = True

playersList = sheet1.get_all_values()[1:]  # 0 - timestamp  1 - name  2 - nickname  3 - game
time.sleep(1)

print("Names loaded")

removeEmpty(playersList)
removeDuplicate(playersList)

playersLOL = []
playersCS = []
prepareGames(playersList)

print("Players for LoL: ", len(playersLOL))
print("Players for CS: ", len(playersCS))

removeExcessivePlayers(playersLOL)
removeExcessivePlayers(playersCS)

print("Players for LoL after cuts: ", len(playersLOL))
print("Players for CS after cuts: ", len(playersCS))

teamsLOL = int(len(playersLOL) / 5)
teamsCS = int(len(playersCS) / 5)

#define rooms
roomsList = []
for i in range(0, teamsLOL):
    roomsList.append([])
for i in range(0, teamsCS):
    roomsList.append([])

#play sound
winsound.PlaySound(None, winsound.SND_PURGE)
winsound.PlaySound('sounds/losu losu losu.wav', winsound.SND_ASYNC)

#choose amount of games
generateGamev2(playersLOL, teamsLOL, "LOL")
generateGamev2(playersLOL, teamsLOL, "LOL")
generateGamev2(playersLOL, teamsLOL, "LOL")
generateGamev2(playersLOL, teamsLOL, "LOL")

#play sound
winsound.PlaySound(None, winsound.SND_PURGE)
winsound.PlaySound('sounds/losu losu losu.wav', winsound.SND_ASYNC)

generateGamev2(playersCS, teamsCS, "CS")
generateGamev2(playersCS, teamsCS, "CS")
generateGamev2(playersCS, teamsCS, "CS")
generateGamev2(playersCS, teamsCS, "CS")

#play sound
winsound.PlaySound(None, winsound.SND_PURGE)
winsound.PlaySound('sounds/oj trudne trudne.wav', winsound.SND_ASYNC)
time.sleep(3)

writeRoomsDOCX(roomsList)

print("Finished successfully")
