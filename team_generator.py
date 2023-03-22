from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtGui import QIcon
from google.oauth2.service_account import Credentials
import gspread
from gspread import utils
import time
import random
import numpy as np
import winsound
import sys
import os


from qt_material import apply_stylesheet
#https://github.com/UN-GCPDS/qt-material#install
#pip install qt-material
#from qt_material import list_themes

#on pycharm: View ->  Tool Windows -> Python Packages -> python-docx -> install
#pip install python-docx
from docx import Document

#on pycharm: View ->  Tool Windows -> Python Packages -> gspread-formatting -> install
#pip install gspread-formatting
from gspread_formatting import *

useName = False
playersList = []
playersLOL = []
playersCS = []
playersValorant = []
playersMK = []
playersFIFA = []
teamsLOL = 0
teamsCS = 0
teamsValorant = 0
roomsList = []
game_number = 1
team_number = 1


def resource_path(relative_path):
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


# set cell color in multiple cells
def setColor(worksheet, color, rowStart, columnStart, rowEnd, columnEnd):
    cellRange = utils.rowcol_to_a1(rowStart, columnStart) + ":" + utils.rowcol_to_a1(rowEnd, columnEnd)
    if color == "red":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 1.0,
                "green": 0.5,
                "blue": 0.5
            }
        })
    elif color == "blue":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 0.5,
                "green": 0.5,
                "blue": 1.0
            }
        })
    elif color == "white":
        worksheet.format(cellRange, {
            "backgroundColor": {
                "red": 1.0,
                "green": 1.0,
                "blue": 1.0
            }
        })


#remove players from the list if name or nickname is missing
def removeEmpty(players):
    i = 0
    while i < len(players):
        for j in range(1, 2):
            if players[i][j] == "":
                del players[i]
                i = i - 1
                continue
        i = i + 1


#remove duplicate entries of players in the list
def removeDuplicate(players):
    i = 0
    while i < len(players):
        j = i + 1
        while j < len(players):
            if players[i][1] == players[j][1] or players[i][2] == players[j][2]:
                del players[j]
                j = j - 1
            j = j + 1
        i = i + 1


#divide players between LOL and CS
def dividePlayers(players):
    for player in players:
        if "lol" in player[3].lower() or "league of legends" in player[3].lower():
            if useName:
                playersLOL.append(player[1])
            else:
                playersLOL.append(player[2])
        elif "cs:go" in player[3].lower() or "csgo" in player[3].lower() or "counter-strike" in player[3].lower() \
                or "counter strike" in player[3].lower():
            if useName:
                playersCS.append(player[1])
            else:
                playersCS.append(player[2])
        elif "valorant" in player[3].lower():
            if useName:
                playersValorant.append(player[1])
            else:
                playersValorant.append(player[2])
        if "mortal kombat" in player[4].lower() or "mk" in player[4].lower():
            if useName:
                playersMK.append(player[1])
            else:
                playersMK.append(player[2])
        elif "fifa" in player[4].lower():
            if useName:
                playersFIFA.append(player[1])
            else:
                playersFIFA.append(player[2])


#remove players if there are too many, for example: reduce from 19 to 10
def removeExcessivePlayers(players, game):
    if game == "LOL" or game == "CS" or game == "Valorant":
        if len(players) < 10:
            print("Impossible to form a team with ", len(players))
            return
        else:
            while len(players) % 10 != 0:
                del players[-1]
    elif game == "MK" or game == "FIFA":
        if len(players) < 8:
            print("Impossible to form bracket with ", len(players))
            return
        else:
            while len(players) % 8 != 0:
                del players[-1]


#Find number of the previous game
def findPrevious(worksheet, text):
    previousGameNumber = 0
    query = re.compile(text + " .*")
    cells = worksheet.findall(query)
    time.sleep(1)
    if len(cells) != 0:
        games = []
        for cell in cells:
            games.append(cell.value)
        games = sorted(games)
        previousGame = games[-1]
        previousGameNumber = int(previousGame[-1])
    return previousGameNumber


#find 1st empty column to write in
def findAvailableColumn(worksheet):
    column = 1
    while column <= 26:
        cell = worksheet.cell(1, column).value
        time.sleep(1)
        if cell is None:
            break
        column = column + 1
    return column


#find 1st empty row to write in
def findAvailableRow(worksheet):
    str_list = list(filter(None, worksheet.col_values(1)))
    time.sleep(1)
    return str(len(str_list)+1)


def generateGame(players, teamNumbers, game):
    global game_number, team_number
    random.shuffle(players)
    if game == "LOL":
        sheet = sheet2
    elif game == "CS":
        sheet = sheet3
    elif game == "Valorant":
        sheet = sheet4

    row = 1
    column = findAvailableColumn(sheet)

    for i in range(0, teamNumbers):
        array = np.empty(shape=[0, 1])
        if i % 2 == 0:
            array = np.append(array, [["GAME " + str(game_number)]], axis=0)
            game_number += 1
        array = np.append(array, [["TEAM " + str(team_number)]], axis=0)
        team_number += 1

        j = i * 5
        while j < i * 5 + 5:
            if game == "LOL":
                roomsList[i].append(players[j])
            elif game == "CS":
                roomsList[i+teamsLOL].append(players[j])
            elif game == "Valorant":
                roomsList[i+teamsLOL+teamsCS].append(players[j])
            array = np.append(array, [[players[j]]], axis=0)
            j = j + 1

        cellRange = utils.rowcol_to_a1(row, column)
        row = row + len(array.tolist())
        sheet.update(cellRange, array.tolist())
        time.sleep(1)

        if i % 2 == 0:
            setColor(sheet, "red", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        else:
            setColor(sheet, "blue", j - 2 + i + i / 2, column, j + 2 + i + i / 2, column)
        time.sleep(1)
        if i % 2 != 0:
            print(game, "game", game_number-1, "generated successfully")


def generateSecondaryGame(players, game):
    global game_number
    random.shuffle(players)
    if game == "MK":
        sheet = sheet5
    elif game == "FIFA":
        sheet = sheet6

    row = 1
    column = findAvailableColumn(sheet)

    for i in range(0, int(len(players)/2)):
        array = np.empty(shape=[0, 1])
        array = np.append(array, [["GAME " + str(game_number)]], axis=0)
        game_number += 1
        array = np.append(array, [[players[2*i]]], axis=0)
        array = np.append(array, [[players[2*i+1]]], axis=0)

        cellRange = utils.rowcol_to_a1(row, column)
        row = row + len(array.tolist())
        sheet.update(cellRange, array.tolist())
        time.sleep(1)

        setColor(sheet, "red", 2*i+i+2, column, 2*i+i+3, column)
        time.sleep(1)

        print(game, "game", game_number-1, "generated successfully")


def openSpreadsheet():
    global spreadsheet, sheet1, sheet2, sheet3, sheet4, sheet5, sheet6
    scope = ['https://www.googleapis.com/auth/spreadsheets',
             'https://www.googleapis.com/auth/drive.file',
             "https://www.googleapis.com/auth/drive"]
    #creds = Credentials.from_service_account_file('genroom-2654a6209a4a.json', scopes=scope)
    #creds = Credentials.from_service_account_file('rage-quit-gaming-539c26e7038f.json', scopes=scope)
    creds = Credentials.from_service_account_file('rqgaming-automation-c71eba04f7ab.json', scopes=scope)
    client = gspread.authorize(creds)

    spreadsheet = client.open("team generator")
    sheet1 = spreadsheet.get_worksheet(0)
    sheet2 = spreadsheet.get_worksheet(1)
    sheet3 = spreadsheet.get_worksheet(2)
    sheet4 = spreadsheet.get_worksheet(3)
    sheet5 = spreadsheet.get_worksheet(4)
    sheet6 = spreadsheet.get_worksheet(5)


class Ui_mainWindow(object):
    def setupUi(self, mainWindow):
        mainWindow.setObjectName("mainWindow")
        mainWindow.resize(866, 900)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Preferred)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(mainWindow.sizePolicy().hasHeightForWidth())
        mainWindow.setSizePolicy(sizePolicy)
        mainWindow.setMinimumSize(QtCore.QSize(866, 900))
        mainWindow.setMaximumSize(QtCore.QSize(866, 900))
        font = QtGui.QFont()
        font.setPointSize(10)
        mainWindow.setFont(font)
        self.centralwidget = QtWidgets.QWidget(mainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.generateButton = QtWidgets.QPushButton(self.centralwidget)
        self.generateButton.setGeometry(QtCore.QRect(320, 640, 231, 71))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.generateButton.setFont(font)
        self.generateButton.setObjectName("generateButton")
        self.gameLabel = QtWidgets.QLabel(self.centralwidget)
        self.gameLabel.setGeometry(QtCore.QRect(30, 0, 131, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameLabel.setFont(font)
        self.gameLabel.setObjectName("gameLabel")
        self.playersLabel = QtWidgets.QLabel(self.centralwidget)
        self.playersLabel.setGeometry(QtCore.QRect(20, 350, 121, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.playersLabel.setFont(font)
        self.playersLabel.setObjectName("playersLabel")
        self.gameSpin1 = QtWidgets.QSpinBox(self.centralwidget)
        self.gameSpin1.setGeometry(QtCore.QRect(400, 40, 81, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameSpin1.setFont(font)
        self.gameSpin1.setMaximum(10)
        self.gameSpin1.setProperty("value", 3)
        self.gameSpin1.setObjectName("gameSpin1")
        self.gameSpin2 = QtWidgets.QSpinBox(self.centralwidget)
        self.gameSpin2.setGeometry(QtCore.QRect(400, 100, 81, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameSpin2.setFont(font)
        self.gameSpin2.setMaximum(10)
        self.gameSpin2.setProperty("value", 3)
        self.gameSpin2.setObjectName("gameSpin2")
        self.gameSpin4 = QtWidgets.QSpinBox(self.centralwidget)
        self.gameSpin4.setGeometry(QtCore.QRect(400, 220, 81, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameSpin4.setFont(font)
        self.gameSpin4.setMaximum(10)
        self.gameSpin4.setProperty("value", 1)
        self.gameSpin4.setObjectName("gameSpin4")
        self.clearButton = QtWidgets.QPushButton(self.centralwidget)
        self.clearButton.setGeometry(QtCore.QRect(610, 690, 221, 71))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.clearButton.setFont(font)
        self.clearButton.setObjectName("clearButton")
        self.gameCheck1 = QtWidgets.QCheckBox(self.centralwidget)
        self.gameCheck1.setGeometry(QtCore.QRect(30, 40, 231, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameCheck1.setFont(font)
        self.gameCheck1.setChecked(True)
        self.gameCheck1.setObjectName("gameCheck1")
        self.gameCheck2 = QtWidgets.QCheckBox(self.centralwidget)
        self.gameCheck2.setGeometry(QtCore.QRect(30, 100, 281, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameCheck2.setFont(font)
        self.gameCheck2.setChecked(True)
        self.gameCheck2.setObjectName("gameCheck2")
        self.gameCheck4 = QtWidgets.QCheckBox(self.centralwidget)
        self.gameCheck4.setGeometry(QtCore.QRect(30, 220, 231, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameCheck4.setFont(font)
        self.gameCheck4.setChecked(True)
        self.gameCheck4.setObjectName("gameCheck4")
        self.playersCheck1 = QtWidgets.QCheckBox(self.centralwidget)
        self.playersCheck1.setGeometry(QtCore.QRect(30, 450, 221, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.playersCheck1.setFont(font)
        self.playersCheck1.setChecked(True)
        self.playersCheck1.setObjectName("playersCheck1")
        self.playersCheck2 = QtWidgets.QCheckBox(self.centralwidget)
        self.playersCheck2.setGeometry(QtCore.QRect(30, 500, 241, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.playersCheck2.setFont(font)
        self.playersCheck2.setChecked(True)
        self.playersCheck2.setObjectName("playersCheck2")
        self.soundCheck = QtWidgets.QCheckBox(self.centralwidget)
        self.soundCheck.setGeometry(QtCore.QRect(710, 20, 111, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.soundCheck.setFont(font)
        self.soundCheck.setChecked(False)
        self.soundCheck.setObjectName("soundCheck")
        self.gamePlayers1 = QtWidgets.QLabel(self.centralwidget)
        self.gamePlayers1.setGeometry(QtCore.QRect(500, 40, 121, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gamePlayers1.setFont(font)
        self.gamePlayers1.setObjectName("gamePlayers1")
        self.label1 = QtWidgets.QLabel(self.centralwidget)
        self.label1.setGeometry(QtCore.QRect(320, 40, 71, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label1.setFont(font)
        self.label1.setObjectName("label1")
        self.gamePlayers2 = QtWidgets.QLabel(self.centralwidget)
        self.gamePlayers2.setGeometry(QtCore.QRect(500, 100, 121, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gamePlayers2.setFont(font)
        self.gamePlayers2.setObjectName("gamePlayers2")
        self.gamePlayers4 = QtWidgets.QLabel(self.centralwidget)
        self.gamePlayers4.setGeometry(QtCore.QRect(500, 220, 121, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gamePlayers4.setFont(font)
        self.gamePlayers4.setObjectName("gamePlayers4")
        self.label2 = QtWidgets.QLabel(self.centralwidget)
        self.label2.setGeometry(QtCore.QRect(320, 100, 71, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label2.setFont(font)
        self.label2.setObjectName("label2")
        self.label4 = QtWidgets.QLabel(self.centralwidget)
        self.label4.setGeometry(QtCore.QRect(320, 220, 71, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label4.setFont(font)
        self.label4.setObjectName("label4")
        self.loadButton = QtWidgets.QPushButton(self.centralwidget)
        self.loadButton.setGeometry(QtCore.QRect(30, 640, 231, 71))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.loadButton.setFont(font)
        self.loadButton.setObjectName("loadButton")
        self.playersRadio1 = QtWidgets.QRadioButton(self.centralwidget)
        self.playersRadio1.setGeometry(QtCore.QRect(30, 400, 141, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.playersRadio1.setFont(font)
        self.playersRadio1.setChecked(False)
        self.playersRadio1.setObjectName("playersRadio1")
        self.playersRadio2 = QtWidgets.QRadioButton(self.centralwidget)
        self.playersRadio2.setGeometry(QtCore.QRect(180, 400, 181, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.playersRadio2.setFont(font)
        self.playersRadio2.setChecked(True)
        self.playersRadio2.setObjectName("playersRadio2")
        self.roomsCheck = QtWidgets.QCheckBox(self.centralwidget)
        self.roomsCheck.setGeometry(QtCore.QRect(30, 550, 241, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.roomsCheck.setFont(font)
        self.roomsCheck.setChecked(True)
        self.roomsCheck.setObjectName("roomsCheck")
        self.nameEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.nameEdit.setGeometry(QtCore.QRect(280, 540, 161, 51))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.nameEdit.setFont(font)
        self.nameEdit.setObjectName("nameEdit")
        self.statusLabel = QtWidgets.QLabel(self.centralwidget)
        self.statusLabel.setGeometry(QtCore.QRect(250, 770, 591, 61))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.statusLabel.setFont(font)
        self.statusLabel.setObjectName("statusLabel")
        self.gameCheck3 = QtWidgets.QCheckBox(self.centralwidget)
        self.gameCheck3.setEnabled(True)
        self.gameCheck3.setGeometry(QtCore.QRect(30, 160, 231, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameCheck3.setFont(font)
        self.gameCheck3.setChecked(True)
        self.gameCheck3.setObjectName("gameCheck3")
        self.label3 = QtWidgets.QLabel(self.centralwidget)
        self.label3.setGeometry(QtCore.QRect(320, 160, 71, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label3.setFont(font)
        self.label3.setObjectName("label3")
        self.gameSpin3 = QtWidgets.QSpinBox(self.centralwidget)
        self.gameSpin3.setGeometry(QtCore.QRect(400, 160, 81, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameSpin3.setFont(font)
        self.gameSpin3.setMaximum(10)
        self.gameSpin3.setProperty("value", 3)
        self.gameSpin3.setObjectName("gameSpin3")
        self.gamePlayers3 = QtWidgets.QLabel(self.centralwidget)
        self.gamePlayers3.setGeometry(QtCore.QRect(500, 160, 121, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gamePlayers3.setFont(font)
        self.gamePlayers3.setObjectName("gamePlayers3")
        self.gameCheck5 = QtWidgets.QCheckBox(self.centralwidget)
        self.gameCheck5.setGeometry(QtCore.QRect(30, 280, 231, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameCheck5.setFont(font)
        self.gameCheck5.setChecked(True)
        self.gameCheck5.setObjectName("gameCheck5")
        self.label5 = QtWidgets.QLabel(self.centralwidget)
        self.label5.setGeometry(QtCore.QRect(320, 280, 71, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label5.setFont(font)
        self.label5.setObjectName("label5")
        self.gameSpin5 = QtWidgets.QSpinBox(self.centralwidget)
        self.gameSpin5.setGeometry(QtCore.QRect(400, 280, 81, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gameSpin5.setFont(font)
        self.gameSpin5.setMaximum(10)
        self.gameSpin5.setProperty("value", 1)
        self.gameSpin5.setObjectName("gameSpin5")
        self.gamePlayers5 = QtWidgets.QLabel(self.centralwidget)
        self.gamePlayers5.setGeometry(QtCore.QRect(500, 280, 121, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.gamePlayers5.setFont(font)
        self.gamePlayers5.setObjectName("gamePlayers5")
        self.prepareButton = QtWidgets.QPushButton(self.centralwidget)
        self.prepareButton.setGeometry(QtCore.QRect(610, 590, 221, 71))
        font = QtGui.QFont()
        font.setPointSize(14)
        self.prepareButton.setFont(font)
        self.prepareButton.setObjectName("prepareButton")
        mainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(mainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 866, 29))
        self.menubar.setObjectName("menubar")
        mainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(mainWindow)
        self.statusbar.setObjectName("statusbar")
        mainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(mainWindow)
        QtCore.QMetaObject.connectSlotsByName(mainWindow)

        self.loadButton.clicked.connect(self.loadPlayers)
        self.generateButton.clicked.connect(self.generate)
        self.clearButton.clicked.connect(self.clearWorksheets)
        self.prepareButton.clicked.connect(self.prepareWorksheets)

    def loadPlayers(self):
        global useName, playersList, playersLOL, playersCS, playersValorant, playersMK, playersFIFA,\
            teamsLOL, teamsCS, teamsValorant, roomsList

        self.statusLabel.setText("Status: Loading players...")
        self.statusLabel.adjustSize()
        app.processEvents()

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(resource_path('sounds/losowanko zapraszam.wav'), winsound.SND_ASYNC)

        # flag to use either names or nicknames
        useName = self.playersRadio2.isChecked()

        # 0 - timestamp  1 - full name  2 - nickname  3 - main game  4 - side game
        playersList = sheet1.get_all_values()[1:]
        time.sleep(1)

        print("Names loaded")

        if self.playersCheck1.isChecked():
            removeEmpty(playersList)
        if self.playersCheck2.isChecked():
            removeDuplicate(playersList)

        playersLOL = []
        playersCS = []
        playersValorant = []
        playersMK = []
        playersFIFA = []
        dividePlayers(playersList)


        removeExcessivePlayers(playersLOL, "LOL")
        removeExcessivePlayers(playersCS, "CS")
        removeExcessivePlayers(playersValorant, "Valorant")
        removeExcessivePlayers(playersMK, "MK")
        removeExcessivePlayers(playersFIFA, "FIFA")

        print("Players for LoL after cuts: ", len(playersLOL))
        print("Players for CS after cuts: ", len(playersCS))
        print("Players for Valorant after cuts: ", len(playersValorant))
        print("Players for Mortal Kombat after cuts: ", len(playersMK))
        print("Players for FIFA after cuts: ", len(playersFIFA))

        self.gamePlayers1.setText("Players: " + str(len(playersLOL)))
        self.gamePlayers2.setText("Players: " + str(len(playersCS)))
        self.gamePlayers3.setText("Players: " + str(len(playersValorant)))
        self.gamePlayers4.setText("Players: " + str(len(playersMK)))
        self.gamePlayers5.setText("Players: " + str(len(playersFIFA)))

        teamsLOL = int(len(playersLOL) / 5)
        teamsCS = int(len(playersCS) / 5)
        teamsValorant = int(len(playersValorant) / 5)

        # define rooms
        roomsList = []
        for i in range(0, teamsLOL):
            roomsList.append([])
        for i in range(0, teamsCS):
            roomsList.append([])
        for i in range(0, teamsValorant):
            roomsList.append([])

        time.sleep(2)

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)

        self.statusLabel.setText("Status: Players loaded")
        self.statusLabel.adjustSize()
        app.processEvents()

    def generate(self):
        global game_number, team_number

        self.loadPlayers()

        self.statusLabel.setText("Status: Generating games...")
        self.statusLabel.adjustSize()
        app.processEvents()

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(resource_path('sounds/losu losu losu.wav'), winsound.SND_ASYNC)

        #LOL
        game_number = 1
        team_number = 1
        if self.gameCheck1.isChecked() and len(playersLOL) >= 10:
            spreadsheet.get_worksheet(1).clear()
            setColor(sheet2, "white", 1, 1, 40, 26)
            for i in range(self.gameSpin1.value()):
                generateGame(playersLOL, teamsLOL, "LOL")

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)
            winsound.PlaySound('sounds/losu losu losu.wav', winsound.SND_ASYNC)

        #CS
        game_number = 1
        team_number = 1
        if self.gameCheck2.isChecked() and len(playersCS) >= 10:
            spreadsheet.get_worksheet(2).clear()
            setColor(sheet3, "white", 1, 1, 40, 26)
            for i in range(self.gameSpin2.value()):
                generateGame(playersCS, teamsCS, "CS")

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)
            winsound.PlaySound(resource_path('sounds/losu losu losu.wav'), winsound.SND_ASYNC)

        # Valorant
        game_number = 1
        team_number = 1
        if self.gameCheck3.isChecked() and len(playersValorant) >= 10:
            spreadsheet.get_worksheet(3).clear()
            setColor(sheet4, "white", 1, 1, 40, 26)
            for i in range(self.gameSpin3.value()):
                generateGame(playersValorant, teamsValorant, "Valorant")

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)
            winsound.PlaySound(resource_path('sounds/losu losu losu.wav'), winsound.SND_ASYNC)

        #Mortal Kombat
        game_number = 1
        team_number = 1
        if self.gameCheck4.isChecked() and len(playersMK) >= 8:
            spreadsheet.get_worksheet(4).clear()
            setColor(sheet5, "white", 1, 1, 40, 26)
            for i in range(self.gameSpin4.value()):
                generateSecondaryGame(playersMK, "MK")

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)
            winsound.PlaySound(resource_path('sounds/losu losu losu.wav'), winsound.SND_ASYNC)

        #FIFA
        game_number = 1
        team_number = 1
        if self.gameCheck5.isChecked() and len(playersFIFA) >= 8:
            spreadsheet.get_worksheet(5).clear()
            setColor(sheet6, "white", 1, 1, 40, 26)
            for i in range(self.gameSpin5.value()):
                generateSecondaryGame(playersFIFA, "FIFA")

        self.statusLabel.setText("Status: Generating completed")
        self.statusLabel.adjustSize()
        app.processEvents()

        # play sound
        if self.soundCheck.isChecked():
            winsound.PlaySound(None, winsound.SND_PURGE)
            winsound.PlaySound(resource_path('sounds/oj trudne trudne.wav'), winsound.SND_ASYNC)
            time.sleep(3)
            winsound.PlaySound(None, winsound.SND_PURGE)

        if self.roomsCheck.isChecked():
            self.writeRoomsDOCX(roomsList)

        print("Finished successfully")

    def prepareWorksheets(self):
        self.statusLabel.setText("Status: Preparing worksheets...")
        self.statusLabel.adjustSize()
        app.processEvents()

        # clear worksheets
        spreadsheet.get_worksheet(1).clear()
        spreadsheet.get_worksheet(2).clear()
        spreadsheet.get_worksheet(3).clear()
        spreadsheet.get_worksheet(4).clear()
        spreadsheet.get_worksheet(5).clear()

        # set column width
        set_column_width(sheet2, 'A:Z', 135)
        set_column_width(sheet3, 'A:Z', 135)
        set_column_width(sheet4, 'A:Z', 135)
        set_column_width(sheet5, 'A:Z', 135)
        set_column_width(sheet6, 'A:Z', 135)

        # clear colors
        setColor(sheet2, "white", 1, 1, 40, 26)
        setColor(sheet3, "white", 1, 1, 40, 26)
        setColor(sheet4, "white", 1, 1, 40, 26)
        setColor(sheet5, "white", 1, 1, 40, 26)
        setColor(sheet6, "white", 1, 1, 40, 26)

        # set font size and bold
        sheet2.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet2.format('A2:Z2', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet2.format('A8:Z8', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet2.format('A14:Z14', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet2.format('A15:Z15', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet2.format('A21:Z21', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet2.format('A27:Z27', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet2.format('A28:Z28', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet2.format('A34:Z34', {'textFormat': {'bold': True, "fontSize": 12}})  # Team

        sheet3.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet3.format('A2:Z2', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet3.format('A8:Z8', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet3.format('A14:Z14', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet3.format('A15:Z15', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet3.format('A21:Z21', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet3.format('A27:Z27', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet3.format('A28:Z28', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet3.format('A34:Z34', {'textFormat': {'bold': True, "fontSize": 12}})  # Team

        sheet4.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet4.format('A2:Z2', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet4.format('A8:Z8', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet4.format('A14:Z14', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet4.format('A15:Z15', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet4.format('A21:Z21', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet4.format('A27:Z27', {'textFormat': {'bold': True, "fontSize": 14}})  # Game
        sheet4.format('A28:Z28', {'textFormat': {'bold': True, "fontSize": 12}})  # Team
        sheet4.format('A34:Z34', {'textFormat': {'bold': True, "fontSize": 12}})  # Team

        sheet5.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A4:Z4', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A7:Z7', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A10:Z10', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A13:Z13', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A16:Z16', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A19:Z19', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet5.format('A22:Z22', {'textFormat': {'bold': True, "fontSize": 12}})  # Game

        sheet6.format('A1:Z1', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A4:Z4', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A7:Z7', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A10:Z10', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A13:Z13', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A16:Z16', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A19:Z19', {'textFormat': {'bold': True, "fontSize": 12}})  # Game
        sheet6.format('A22:Z22', {'textFormat': {'bold': True, "fontSize": 12}})  # Game

        # this is the amount of previous operations
        time.sleep(58)

        print("Worksheets prepared")
        self.statusLabel.setText("Status: Worksheets prepared")
        self.statusLabel.adjustSize()
        app.processEvents()

    def clearWorksheets(self):
        self.statusLabel.setText("Status: Clearing worksheets...")
        self.statusLabel.adjustSize()
        app.processEvents()

        # clear worksheets
        spreadsheet.get_worksheet(1).clear()
        spreadsheet.get_worksheet(2).clear()
        spreadsheet.get_worksheet(3).clear()
        spreadsheet.get_worksheet(4).clear()
        spreadsheet.get_worksheet(5).clear()

        # clear colors
        setColor(sheet2, "white", 1, 1, 40, 26)
        setColor(sheet3, "white", 1, 1, 40, 26)
        setColor(sheet4, "white", 1, 1, 40, 26)
        setColor(sheet5, "white", 1, 1, 40, 26)
        setColor(sheet6, "white", 1, 1, 40, 26)

        # this is the amount of previous operations
        #time.sleep(10)

        self.statusLabel.setText("Status: Worksheets cleared")
        self.statusLabel.adjustSize()
        app.processEvents()

    def writeRoomsDOCX(self, rooms):
        document = Document()
        document.add_heading('Rooms', 0)
        for i in range(0, len(rooms)):
            heading = document.add_heading('Room' + (str(i + 1)))
            if i < teamsLOL:
                heading.add_run(' (LOL)')
            elif i < teamsLOL + teamsCS:
                heading.add_run(' (CS)')
            else:
                heading.add_run(' (Valorant)')
            for j in range(0, len(rooms[i])):
                if j % 5 == 0:
                    paragraph = document.add_paragraph('Game ' + str(int(j / 5 + 1)) + ':  ')
                paragraph.add_run(rooms[i][j])
                if (j + 1) % 5 != 0:
                    paragraph.add_run(', ')
        document.save(self.nameEdit.text() + '.docx')

        print("Rooms distribution file DOCX generated successfully")

    def writeRoomsTXT(self, rooms):
        f = open("rooms.txt", "w")
        for i in range(0, len(rooms)):
            if i != 0:
                f.write("\n\n")
            f.write("Room " + str(i + 1))

            if i < teamsLOL:
                f.write(" (LOL)")
            elif i < teamsLOL + teamsCS:
                f.write("(CS)")
            else:
                f.write("(Valorant)")

            for j in range(0, len(rooms[i])):
                if j % 5 == 0:
                    f.write("\nGame " + str(int(j / 5 + 1)) + ":  ")
                f.write(rooms[i][j])
                if (j + 1) % 5 != 0:
                    f.write(", ")
        f.close()
        print("Rooms distribution file generated successfully")

    def retranslateUi(self, mainWindow):
        _translate = QtCore.QCoreApplication.translate
        mainWindow.setWindowTitle(_translate("mainWindow", "Team Generator"))
        self.generateButton.setText(_translate("mainWindow", "Generate"))
        self.gameLabel.setText(_translate("mainWindow", "Game Options"))
        self.playersLabel.setText(_translate("mainWindow", "Player options"))
        self.clearButton.setText(_translate("mainWindow", "Clear sheets"))
        self.gameCheck1.setText(_translate("mainWindow", "League of Legends"))
        self.gameCheck2.setText(_translate("mainWindow", "Counter-Strike: Global Offensive"))
        self.gameCheck4.setText(_translate("mainWindow", "Mortal Kombat"))
        self.playersCheck1.setText(_translate("mainWindow", "Remove empty players"))
        self.playersCheck2.setText(_translate("mainWindow", "Remove duplicate players"))
        self.soundCheck.setText(_translate("mainWindow", "Sounds"))
        self.gamePlayers1.setText(_translate("mainWindow", "Players: "))
        self.label1.setText(_translate("mainWindow", "Games: "))
        self.gamePlayers2.setText(_translate("mainWindow", "Players: "))
        self.gamePlayers4.setText(_translate("mainWindow", "Players: "))
        self.label2.setText(_translate("mainWindow", "Games: "))
        self.label4.setText(_translate("mainWindow", "Games: "))
        self.loadButton.setText(_translate("mainWindow", "Load players"))
        self.playersRadio1.setText(_translate("mainWindow", "Nickname"))
        self.playersRadio2.setText(_translate("mainWindow", "Full name"))
        self.roomsCheck.setText(_translate("mainWindow", "Generate rooms docx file:"))
        self.nameEdit.setText(_translate("mainWindow", "rooms"))
        self.statusLabel.setText(_translate("mainWindow", "Status: Waiting..."))
        self.gameCheck3.setText(_translate("mainWindow", "Valorant"))
        self.label3.setText(_translate("mainWindow", "Games: "))
        self.gamePlayers3.setText(_translate("mainWindow", "Players: "))
        self.gameCheck5.setText(_translate("mainWindow", "Fifa"))
        self.label5.setText(_translate("mainWindow", "Games: "))
        self.gamePlayers5.setText(_translate("mainWindow", "Players: "))
        self.prepareButton.setText(_translate("mainWindow", "Prepare sheets"))


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)

    # icon
    app_icon = QIcon(resource_path("images/icon.png"))
    app.setWindowIcon(app_icon)

    # styles
    apply_stylesheet(app, theme='dark_lightgreen.xml')

    mainWindow = QtWidgets.QMainWindow()
    ui = Ui_mainWindow()
    ui.setupUi(mainWindow)
    openSpreadsheet()
    mainWindow.show()
    sys.exit(app.exec_())

